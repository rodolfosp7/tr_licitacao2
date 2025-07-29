import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os

st.set_page_config(page_title="Agente de Licitações - Termo de Referência", layout="wide")

st.title("Agente de Licitações - Geração de Termo de Referência")

# Entrada do objeto
objeto = st.text_input("Informe o OBJETO da contratação (ex: fornecimento de gás de cozinha):")

# Caminhos das imagens de cabeçalho e rodapé
logo_path = "logo-prefeitura.png"
rodape_path = "rodapé.png"

if objeto:
    st.markdown("### Estrutura do Documento")
    st.markdown("Preencheremos cada item com informações detalhadas, conforme modelo.")

    if st.button("Gerar Documento Word"):
        doc = Document()
        section = doc.sections[0]

        # Cabeçalho com logo
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run()
        if os.path.exists(logo_path):
            run.add_picture(logo_path, width=Inches(2.5))

        # Corpo do documento
        doc.add_heading("1. DAS CONDIÇÕES GERAIS DA CONTRATAÇÃO", level=1)
        doc.add_paragraph(f"O presente termo de referência tem por objeto: {objeto}. Todas as especificações deverão seguir a legislação vigente.")

        doc.add_heading("2. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO E FUNDAMENTAÇÃO LEGAL", level=1)
        doc.add_paragraph("Item detalhado com justificativas completas conforme a necessidade do objeto.")

        doc.add_heading("3. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO CONSIDERADO O CICLO DE VIDA DO OBJETO E ESPECIFICAÇÃO DOS SERVIÇOS", level=1)
        doc.add_paragraph("Todas as soluções avaliadas e a solução recomendada com ciclo de vida.")

        doc.add_heading("4. REQUISITOS DA CONTRATAÇÃO", level=1)
        doc.add_paragraph("Lista de requisitos técnicos e legais.")

        doc.add_heading("5. MODELO DE EXECUÇÃO CONTRATUAL", level=1)
        doc.add_paragraph("Descrição de execução, fiscalização e sanções.")

        doc.add_heading("6. CRITÉRIOS DE MEDIÇÃO", level=1)
        doc.add_paragraph("Definir detalhadamente como será feita a medição do objeto.")

        # Rodapé com imagem
        footer = section.footer
        footer_para = footer.paragraphs[0]
        run_footer = footer_para.add_run()
        if os.path.exists(rodape_path):
            run_footer.add_picture(rodape_path, width=Inches(6))

        # Exportar para buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📄 Baixar Termo de Referência em Word",
            data=buffer,
            file_name=f"Termo_Referencia_{objeto}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
