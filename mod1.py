import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os
from typing import List, Tuple, Dict
from docx.oxml.ns import qn

st.set_page_config(page_title="Agente de Licitações - Termo de Referência", layout="wide")

st.title("Agente de Licitações – TR detalhado (com modelo interno + importação de Word)")

# =============================
# Campos-base (dados comuns)
# =============================
with st.sidebar:
    st.header("⚙️ Parâmetros Gerais")
    municipio = st.text_input("Município", value="Brasnorte-MT")
    secretaria = st.text_input("Setor Requisitante", value="Secretaria Municipal de Administração")
    objeto = st.text_input("OBJETO (ex.: fornecimento de gás de cozinha)", value="")
    modalidade = st.selectbox(
        "Modalidade",
        [
            "Pregão Eletrônico",
            "Concorrência",
            "Dispensa",
            "Inexigibilidade",
            "Credenciamento"
        ], index=0
    )
    usa_srp = st.checkbox("Sistema de Registro de Preços (SRP)", value=True)
    criterio = st.selectbox("Critério de Julgamento", ["Menor preço por item", "Menor preço global", "Maior desconto", "Técnica e preço"], index=0)
    prazo_vigencia_meses = st.number_input("Vigência (meses)", min_value=1, value=12)
    permite_prorrogacao = st.checkbox("Permite prorrogação (arts. 106 e 107 da Lei 14.133/2021)", value=True)
    decreto_luxo = st.text_input("Decreto municipal (vedação a bem de luxo)", value="03/2024")

# Caminhos das imagens de cabeçalho e rodapé (pré-carregadas no ambiente)
LOGO_PATH = "/mnt/data/logo-prefeitura.png"
RODAPE_PATH = "/mnt/data/rodapé.png"

# ======================================
# Funções utilitárias para importar .docx
# ======================================

def ler_modelo_docx(file) -> List[Tuple[str, str]]:
    """Lê um arquivo .docx de modelo e retorna uma lista de (heading, texto_acumulado).
    Considera como heading qualquer parágrafo com estilo 'Heading X'.
    """
    doc = Document(file)
    blocos = []
    titulo_atual = None
    buffer = []
    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "")
        texto = p.text.strip()
        if style_name and style_name.lower().startswith("heading") and texto:
            # Fecha bloco anterior
            if titulo_atual is not None:
                blocos.append((titulo_atual, "
".join(buffer).strip()))
                buffer = []
            titulo_atual = texto
        else:
            if texto:
                buffer.append(texto)
    # Último bloco
    if titulo_atual is not None:
        blocos.append((titulo_atual, "
".join(buffer).strip()))
    return blocos


def aplicar_placeholders(texto: str, context: Dict[str, str]) -> str:
    for k, v in context.items():
        texto = texto.replace(f"{{{{{k}}}}}", v)
    return texto

# ======================================
# Template interno detalhado por seções
# ======================================

def template_interno(secao: int, ctx: Dict[str, str]) -> List[Tuple[str, str]]:
    """Retorna uma lista de tuplas (heading, body) para cada seção.
    O conteúdo é detalhado, com subitens e narrativa robusta, baseado no modelo fornecido pelo usuário.
    """
    OBJ = ctx.get("OBJETO", "objeto")
    MUN = ctx.get("MUNICIPIO", "Município")
    MOD = ctx.get("MODALIDADE", "Pregão Eletrônico")
    SRP = ctx.get("SRP", "Sim")
    CRI = ctx.get("CRITERIO", "Menor preço por item")
    VIG = ctx.get("VIGENCIA", "12")
    DECR = ctx.get("DECRETO_LUXO", "03/2024")

    blocks: List[Tuple[str, str]] = []

    if secao == 1:
        heading = "1. DAS CONDIÇÕES GERAIS DA CONTRATAÇÃO"
        body = (
            f"1.1 O presente termo de referência tem por objeto o REGISTRO DE PREÇO PARA FUTURA E EVENTUAL CONTRATAÇÃO DE EMPRESA ESPECIALIZADA EM {OBJ}, com sede localizada no município de {MUN}, em conformidade com as especificações de descrição e quantidade detalhadamente elencadas neste documento, amparada pelas disposições legais vigentes que regulam tal procedimento, visando atender as necessidades da Prefeitura Municipal de {MUN} e de suas Secretarias Municipais.

"
            f"1.2 O objeto desta contratação não se enquadra como sendo de bem de luxo, conforme Decreto Municipal nº {DECR}.

"
            f"1.3 O prazo de vigência da contratação é de {VIG} ( {VIG} ) meses, contados da data de assinatura da ARP (Ata de Registro de Preços) ou do Contrato conforme celebrado, na forma do artigo 105 da Lei nº 14.133/2021, podendo o mesmo ser prorrogado a critério da Administração Pública.

"
            "1.4 O prazo de vigência poderá ser prorrogado, desde que haja interesse de ambas as partes, na forma autorizada pelos artigos 106 e 107, da Lei nº 14.133/2021."
        )
        blocks.append((heading, body))

    if secao == 2:
        heading = "2. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO E FUNDAMENTAÇÃO LEGAL"
        body = (
            f"2.1 A presente contratação se fundamenta na necessidade em possuir {OBJ} para atender as necessidades do Município de {MUN}, em todas as Secretarias Municipais, utilizados no desempenho de suas atividades e cumprimento de sua missão institucional.

"
            "2.2 A demanda se destina ao atendimento de servidores, profissionais, consultores, técnicos, representantes de órgãos públicos, fornecedores, prestadores de serviços e demais colaboradores envolvidos em atividades de interesse público (cursos, oficinas, treinamentos, execuções contratuais, inspeções, auditorias, reuniões técnicas e operacionais).

"
            f"2.3 A contratação justifica-se pelos princípios da eficiência, economicidade e continuidade do serviço público, assegurando condições adequadas de segurança, regularidade, conforto e conformidade legal na execução de {OBJ}.

"
            f"2.4 O procedimento licitatório adotará a modalidade {MOD}{' com utilização do Sistema de Registro de Preços (SRP)' if SRP=='Sim' else ''}, com critério de julgamento '{CRI}', conforme os arts. 6º, 28, 82 e seguintes da Lei nº 14.133/2021 e, quando aplicável, o Decreto Federal nº 11.462/2023 (SRP)."
        )
        blocks.append((heading, body))

    if secao == 3:
        heading = "3. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO CONSIDERADO O CICLO DE VIDA DO OBJETO E ESPECIFICAÇÃO DOS SERVIÇOS"
        body = (
            f"3.1 O objetivo é selecionar a proposta mais vantajosa para {OBJ}, observando requisitos de qualidade, prazos e conformidade regulatória.

"
            f"3.2 Ciclo de vida do objeto: planejamento da demanda; seleção do fornecedor; formalização contratual; execução (fornecimento, logística, conferência, recebimento provisório/definitivo); avaliação de desempenho; e encerramento, com análise de indicadores e lições aprendidas.

"
            "3.3 Alternativas avaliadas:
"
            "• Solução 1 – Execução direta pela Administração: potencial controle direto, porém, em geral, inviável por ausência de equipe técnica, infraestrutura dedicada, riscos operacionais e custos de implantação/manutenção.
"
            "• Solução 2 – Execução indireta (terceirização/fornecedor especializado): transferência de riscos operacionais ao contratado, atendimento a normas técnicas e sanitárias, maior flexibilidade e agilidade, com necessidade de fiscalização permanente pela Administração.

"
            "Conclusão: a Solução 2 mostra-se mais eficiente, econômica e segura, em conformidade com a Lei nº 14.133/2021.

"
            f"3.4 Especificações resumidas do objeto (adaptar conforme {OBJ}):
"
            "• Qualidade e conformidade com normas técnicas aplicáveis;
"
            "• Garantia de fornecimento contínuo;
"
            "• Atendimento a padrões de segurança, saúde e meio ambiente, quando aplicável;
"
            "• Emissão de nota fiscal com detalhamento por item e período;
"
            "• Suporte e atendimento em dias úteis e, quando necessário, fins de semana e feriados."
        )
        blocks.append((heading, body))

    if secao == 4:
        heading = "4. REQUISITOS DA CONTRATAÇÃO"
        body = (
            "4.1 Requisitos legais e habilitação: CNPJ ativo; regularidade fiscal e trabalhista; inscrição em cadastros pertinentes; atendimento à LGPD quando aplicável; atestados de capacidade técnica compatíveis com o objeto; e demais documentos previstos em edital.

"
            f"4.2 Requisitos técnicos mínimos (adaptar ao {OBJ}): conformidade com normas da ABNT/INMETRO e/ou regulatórias; padrões de segurança e qualidade; logística de fornecimento; e comprovação de capacidade operacional para atendimento à demanda.

"
            "4.3 Requisitos funcionais: atendimento sob demanda, sem cota mínima; cumprimento de prazos; suporte adequado; emissão de comprovantes/documentos para fins de controle e fiscalização administrativos.

"
            "4.4 Sustentabilidade (quando aplicável): gestão eficiente de água e energia; produtos e insumos com menor impacto ambiental; destinação adequada de resíduos; acessibilidade e inclusão.

"
            "4.5 Conformidade legal: observância integral da Lei nº 14.133/2021, normas sanitárias, de segurança e ambientais aplicáveis, além de orientações dos órgãos de controle."
        )
        blocks.append((heading, body))

    if secao == 5:
        heading = "5. MODELO DE EXECUÇÃO CONTRATUAL"
        body = (
            "5.1 O contrato deverá ser executado fielmente pelas partes; comunicações preferencialmente por escrito; possibilidade de reunião inicial para apresentação do plano de fiscalização.

"
            "5.2 Fiscalização (art. 117 da Lei nº 14.133/2021): o(s) fiscal(is) acompanharão a execução, registrarão ocorrências, notificarão correções, verificarão manutenção das condições de habilitação, empenho, pagamentos, garantias e eventuais glosas.

"
            "5.3 Gestão do contrato: o gestor consolidará registros formais (ordens de serviço, ocorrências, alterações, prorrogações), avaliará desempenho com base em indicadores e proporá medidas saneadoras quando necessário; elaborará relatório final ao término.

"
            "5.4 Extinção contratual: observar Arts. 137 a 139 da Lei nº 14.133/2021, incluindo hipóteses por inadimplemento, caso fortuito/força maior, razões de interesse público, entre outras; prever consequências e direitos, inclusive devolução de garantia e pagamentos devidos, quando cabível."
        )
        blocks.append((heading, body))

    if secao == 6:
        heading = "6. CRITÉRIOS DE MEDIÇÃO"
        body = (
            f"6.1 A medição será mensal e baseada no serviço/bem efetivamente {('prestado' if 'serviço' in OBJ.lower() else 'fornecido')} e atestado pela Administração.

"
            "6.2 Unidade de medida: conforme item e especificações (ex.: unidade, litro, kg, diária), respeitando ordens de fornecimento/serviço.

"
            "6.3 Documentos de medição: relação detalhada dos itens/quantitativos; relatórios de execução/entrega; notas fiscais compatíveis com preços registrados; comprovação de autorização formal.

"
            "6.4 Conferência e atesto: o gestor/fiscal conferirá informações, atestará relatórios e validará notas para liberação de pagamento, se atendidas as exigências contratuais.

"
            "6.5 Penalidades por divergências: inconsistências sem justificativa poderão ensejar glosas proporcionais, suspensão de pagamento e aplicação de sanções, nos termos da Lei nº 14.133/2021."
        )
        blocks.append((heading, body))

    return blocks

# ======================================
# Construção dos blocos do documento
# ======================================

def construir_blocos(modo: str, ctx: Dict[str, str], modelos_importados: Dict[str, List[Tuple[str, str]]], modelo_escolhido: str) -> List[Tuple[str, str]]:
    """Retorna lista de (heading, body). modo: 'interno' ou 'importado'."""
    if modo == "importado" and modelo_escolhido and modelo_escolhido in modelos_importados:
        blocos_raw = modelos_importados[modelo_escolhido]
        blocos = []
        for h, b in blocos_raw:
            blocos.append((aplicar_placeholders(h, ctx), aplicar_placeholders(b, ctx)))
        return blocos
    # modo interno: monta 1..6
    blocos = []
    for i in range(1, 7):
        blocos.extend(template_interno(i, ctx))
    return blocos

# ======================================
# UI – escolha da fonte do conteúdo
# ======================================
col1, col2 = st.columns([2, 1])
with col1:
    st.subheader("🧩 Fonte do conteúdo do TR")
    fonte = st.radio("Escolha como gerar o conteúdo:", ["Template interno detalhado", "Extrair de modelo Word (upload)"])

with col2:
    uploaded_files = st.file_uploader("Modelos Word (.docx)", type=["docx"], accept_multiple_files=True)

modelos_importados: Dict[str, List[Tuple[str, str]]] = {}
if uploaded_files:
    for f in uploaded_files:
        try:
            modelos_importados[f.name] = ler_modelo_docx(f)
        except Exception as e:
            st.warning(f"Não foi possível ler o modelo: {f.name} ({e})")

modelo_escolhido = None
if fonte == "Extrair de modelo Word (upload)":
    if modelos_importados:
        modelo_escolhido = st.selectbox("Selecione o arquivo-base:", list(modelos_importados.keys()))
    else:
        st.info("Envie ao menos um arquivo .docx com headings.")

# ======================================
# Placeholders (para uso em modelos importados)
# ======================================
ctx = {
    "OBJETO": objeto or "(definir objeto)",
    "MUNICIPIO": municipio,
    "SECRETARIA": secretaria,
    "MODALIDADE": modalidade,
    "SRP": "Sim" if usa_srp else "Não",
    "CRITERIO": criterio,
    "VIGENCIA": str(prazo_vigencia_meses),
    "DECRETO_LUXO": decreto_luxo,
}

# =============================
# Pré-visualização
# =============================
if objeto:
    st.markdown("---")
    st.subheader("👁️ Pré-visualização do conteúdo")
    modo = "importado" if fonte.startswith("Extrair") and modelo_escolhido else "interno"
    blocos = construir_blocos(modo, ctx, modelos_importados, modelo_escolhido)

    with st.expander("Mostrar prévia estruturada", expanded=True):
        for heading, body in blocos:
            st.markdown(f"**{heading}**")
            st.markdown(body.replace("
", "  
"))
            st.markdown("")
else:
    st.info("➡️ Preencha o OBJETO na barra lateral para gerar a prévia e o Word.")

# =============================
# Geração do documento Word
# =============================

def gerar_docx(blocos: List[Tuple[str, str]], ctx: Dict[str, str]) -> BytesIO:
    doc = Document()
    section = doc.sections[0]

    # Cabeçalho com logo
    try:
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run()
        if os.path.exists(LOGO_PATH):
            run.add_picture(LOGO_PATH, width=Inches(2.5))
    except Exception:
        pass

    # Título inicial (opcional)
    doc.add_heading("TERMO DE REFERÊNCIA", level=0)
    p_meta = doc.add_paragraph()
    p_meta.add_run("Município: ").bold = True
    p_meta.add_run(ctx.get("MUNICIPIO", ""))
    p_meta.add_run("  |  Setor requisitante: ").bold = True
    p_meta.add_run(ctx.get("SECRETARIA", ""))

    # Corpo
    for heading, body in blocos:
        level = 1 if heading[:1].isdigit() else 2
        doc.add_heading(heading, level=level)
        for par in body.split("

"):
            doc.add_paragraph(par)

    # Rodapé com imagem
    try:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        run_footer = footer_para.add_run()
        if os.path.exists(RODAPE_PATH):
            run_footer.add_picture(RODAPE_PATH, width=Inches(6))
    except Exception:
        pass

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.markdown("---")
colA, colB = st.columns([1, 2])
with colA:
    gerar = st.button("📄 Gerar Word agora", type="primary", use_container_width=True)
with colB:
    st.caption("O documento conterá cabeçalho e rodapé e todas as seções detalhadas. Você pode usar o template interno ou importar um modelo .docx e usar placeholders como {{OBJETO}}, {{MUNICIPIO}}, {{VIGENCIA}}, etc.")

if gerar:
    if not objeto:
        st.error("Informe o OBJETO na barra lateral antes de gerar o documento.")
    else:
        modo = "importado" if fonte.startswith("Extrair") and modelo_escolhido else "interno"
        blocos = construir_blocos(modo, ctx, modelos_importados, modelo_escolhido)
        if not blocos:
            st.error("Não há conteúdo pronto para gerar. Verifique o objeto ou o modelo importado.")
        else:
            docx_buffer = gerar_docx(blocos, ctx)
            st.download_button(
                label="⬇️ Baixar Termo de Referência (.docx)",
                data=docx_buffer,
                file_name=f"TR_{ctx['OBJETO'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
