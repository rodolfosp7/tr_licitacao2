import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os

st.set_page_config(page_title="Agente de Licitações - Termo de Referência", layout="wide")

st.title("Agente de Licitações - Geração de Termo de Referência")

# Entrada do objeto
objeto = st.text_input("Informe o OBJETO da contratação (ex: fornecimento de gás de cozinha):")

# Caminhos das imagens de cabeçalho e rodapé
logo_path = "logo-prefeitura.png"
rodape_path = "rodapé.png"

def gerar_texto_detalhado(secao, objeto):
    if secao == 1:
        return (f"1.1 O presente termo de referência tem por objeto o REGISTRO DE PREÇO PARA FUTURA E EVENTUAL CONTRATAÇÃO DE EMPRESA ESPECIALIZADA EM {objeto}, em conformidade com as especificações e quantidades descritas neste documento, amparada pelas disposições legais vigentes que regulam tal procedimento, visando atender as necessidades da Administração.\n"
                "1.2 O objeto desta contratação não se enquadra como sendo de bem de luxo, conforme Decreto Municipal nº 03/2024.\n"
                "1.3 O prazo de vigência será de 12 (doze) meses, podendo ser prorrogado conforme artigos 106 e 107 da Lei nº 14.133/2021.")
    if secao == 2:
        return ("2.1 A presente contratação se fundamenta na necessidade da Administração em atender suas demandas com eficiência e economicidade.\n"
                f"2.2 O {objeto} será destinado ao uso contínuo das secretarias municipais e serviços essenciais.\n"
                "2.3 Justifica-se pelo interesse público, pela ausência de estrutura própria da Administração e pela necessidade de segurança, regularidade e conformidade legal.\n"
                "2.4 A licitação será conduzida na modalidade Pregão Eletrônico, SRP, conforme artigos 28 e 82 da Lei nº 14.133/2021 e Decreto Federal nº 11.462/2023.")
    if secao == 3:
        return (f"3.1 O ciclo de vida do {objeto} abrange desde o planejamento da demanda, contratação, fornecimento e utilização final do bem.\n"
                "3.2 Foram avaliadas soluções alternativas e selecionada a mais vantajosa conforme critérios de economicidade e qualidade.")
    if secao == 4:
        return ("4.1 A empresa deverá comprovar capacidade técnica e regularidade fiscal.\n"
                "4.2 Apresentar CNPJ ativo, alvarás e atestados técnicos.\n"
                f"4.3 O {objeto} deverá atender padrões de qualidade, segurança e conformidade.\n"
                "4.4 Produtos fora de especificação não serão aceitos.")
    if secao == 5:
        return ("5.1 O contrato será executado fielmente pelas partes.\n"
                "5.2 A fiscalização será realizada por fiscais designados, que acompanharão a entrega, qualidade e conformidade do objeto.\n"
                "5.3 Ocorrências serão registradas em relatórios, sendo aplicadas sanções quando necessário, conforme arts. 117 e 124 da Lei nº 14.133/2021.")
    if secao == 6:
        return ("6.1 A medição será feita mensalmente com base nos quantitativos efetivamente entregues e aceitos.\n"
                "6.2 Serão aceitos apenas bens previamente autorizados e dentro do escopo contratual.")

if objeto:
    st.markdown("### Estrutura do Documento")
    st.markdown("Cada item será gerado com informações detalhadas e subitens numerados conforme os Termos de Referência modelo.")

    if st.button("Gerar Documento Word"):
        doc = Document()
        section = doc.sections[0]

        # Cabeçalho com logo
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run()
        if os.path.exists(logo_path):
            run.add_picture(logo_path, width=Inches(2.5))

        # Corpo do documento com detalhes e subitens numerados
        doc.add_heading("1. DAS CONDIÇÕES GERAIS DA CONTRATAÇÃO", level=1)
        doc.add_paragraph(gerar_texto_detalhado(1, objeto))

        doc.add_heading("2. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO E FUNDAMENTAÇÃO LEGAL", level=1)
        doc.add_paragraph(gerar_texto_detalhado(2, objeto))

        doc.add_heading("3. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO CONSIDERADO O CICLO DE VIDA DO OBJETO E ESPECIFICAÇÃO DOS SERVIÇOS", level=1)
        doc.add_paragraph(gerar_texto_detalhado(3, objeto))

        doc.add_heading("4. REQUISITOS DA CONTRATAÇÃO", level=1)
        doc.add_paragraph(gerar_texto_detalhado(4, objeto))

        doc.add_heading("5. MODELO DE EXECUÇÃO CONTRATUAL", level=1)
        doc.add_paragraph(gerar_texto_detalhado(5, objeto))

        doc.add_heading("6. CRITÉRIOS DE MEDIÇÃO", level=1)
        doc.add_paragraph(gerar_texto_detalhado(6, objeto))

        # Rodapé com imagem
        footer = section.footer
        footer_para = footer.paragraphs[0]
        run_footer = footer_para.add_run()
        if os.path.exists(rodape_path):
            run_footer.add_picture(rodape_path, width=Inches(6))

        # Exportar para buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📄 Baixar Termo de Referência em Word",
            data=buffer,
            file_name=f"Termo_Referencia_{objeto}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
