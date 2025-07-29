# -*- coding: utf-8 -*-
"""
Módulo: importacao_e_combinacao_tr.py

Objetivo
--------
1) **Ler modelos DOCX** como os enviados (ex.: "TR - SERVIÇOS POSTAIS - CORREIOS.docx"),
   mesmo quando os títulos NÃO estão com estilos de Heading, mas sim com **numeração manual**
   (1., 1.1, 2., 3. etc.).
2) **Construir seções** preservando a ordem dos elementos (parágrafos e tabelas) do DOCX.
3) **Combinar** o conteúdo importado com o **template interno** do app, com três modos:
   - "modelo": usa somente o que veio do DOCX;
   - "template": usa somente o template interno;
   - **"complementar" (recomendado)**: usa o conteúdo do DOCX como base e **insere as seções
     que estiverem faltando** a partir do template interno (respeitando a numeração).
4) **Gerar o DOCX final** (opcionalmente já com cabeçalho e rodapé com imagens fornecidas).

Este arquivo foi escrito para ser facilmente integrado ao Streamlit do seu app. Basta:
- importar as funções no seu `app.py`/`appTR.py`;
- chamar `ler_modelo_docx()` quando o usuário fizer upload do modelo;
- escolher o modo de combinação na UI e chamar `combinar_secoes()`;
- por fim, chamar `gerar_docx()` para baixar o arquivo final.

Requisitos:
    python-docx

Imagens de cabeçalho/rodapé (se desejar):
    /mnt/data/logo-prefeitura.png
    /mnt/data/rodapé.png
"""
from __future__ import annotations
import re
from dataclasses import dataclass, field
from typing import List, Tuple, Dict, Optional

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches

# ==========================
# Estruturas de dados
# ==========================

@dataclass
class Elemento:
    tipo: str  # 'p' ou 'table'
    payload: object  # texto (str) quando 'p'; tabela (python-docx table) quando 'table'

@dataclass
class Secao:
    titulo: str
    numero: Optional[str]  # ex.: '1', '1.1', '2'
    elementos: List[Elemento] = field(default_factory=list)

# ==========================
# Utilitários de parsing
# ==========================

_regex_inicio_secao = re.compile(r"^(?P<num>\d+(?:\.\d+)*)\s*[-–—)]?\s+", re.UNICODE)

def _iter_elementos_em_ordem(doc: Document):
    """Itera parágrafos e tabelas na ordem em que aparecem no corpo do documento.
    Retorna tuplas (tipo, objeto), onde tipo ∈ {"p", "table"}.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield ("p", child)
        elif isinstance(child, CT_Tbl):
            yield ("table", child)


def _wrap_paragraph(doc: Document, ct_p: CT_P):
    # Constrói um objeto Paragraph python-docx a partir do CT_P bruto
    return doc.paragraphs[0]._parent._paragraphs._p_to_paragraph(ct_p)


def _wrap_table(doc: Document, ct_tbl: CT_Tbl):
    # Constrói um objeto Table python-docx a partir do CT_Tbl bruto
    return doc._part._element_to_table(ct_tbl)


def _tenta_numero_secao(texto: str) -> Tuple[Optional[str], Optional[str]]:
    """Se o parágrafo parecer início de seção numérica, retorna (numero, titulo_completo).
    Caso contrário, (None, None).
    """
    raw = texto.strip()
    if not raw:
        return None, None
    m = _regex_inicio_secao.match(raw)
    if m:
        numero = m.group("num")
        return numero, raw
    return None, None

# ==========================
# Leitura de DOCX em seções
# ==========================

def ler_modelo_docx(file_path_or_bytes) -> List[Secao]:
    """Lê um DOCX e segmenta em seções por **numeração manual** (1., 1.1, 2., ...).
    Preserva parágrafos e tabelas na ordem.

    Retorna: lista de Secao, cada uma com `titulo`, `numero` (se detectado) e `elementos`.
    """
    doc = Document(file_path_or_bytes)
    secoes: List[Secao] = []
    secao_atual: Optional[Secao] = None

    # Precisamos recriar objetos python-docx a partir de CT_P/CT_Tbl preservando ordem.
    for tipo, child in _iter_elementos_em_ordem(doc):
        if tipo == "p":
            p = _wrap_paragraph(doc, child)
            texto = p.text.strip()
            numero, titulo = _tenta_numero_secao(texto)
            if numero is not None:
                # Fechar seção anterior
                if secao_atual is not None:
                    secoes.append(secao_atual)
                secao_atual = Secao(titulo=titulo, numero=numero, elementos=[])
            else:
                if secao_atual is None:
                    # Conteúdo prévio sem numeração: cria uma seção 0.
                    secao_atual = Secao(titulo="0. PREÂMBULO", numero="0", elementos=[])
                secao_atual.elementos.append(Elemento("p", texto))
        else:  # table
            tbl = _wrap_table(doc, child)
            if secao_atual is None:
                secao_atual = Secao(titulo="0. PREÂMBULO", numero="0", elementos=[])
            secao_atual.elementos.append(Elemento("table", tbl))

    if secao_atual is not None:
        secoes.append(secao_atual)

    return secoes

# ==========================
# Template interno (exemplo)
# ==========================

def template_interno_padrao() -> List[Secao]:
    """Retorna seções do template interno. Aqui estão as 6 seções básicas
    (ajuste conforme seu template real). Cada seção entra com um parágrafo
    placeholder para que possam ser mescladas.
    """
    def s(num, titulo, texto):
        return Secao(titulo=f"{num} {titulo}", numero=str(num).split()[0], elementos=[Elemento("p", texto)])

    return [
        Secao(titulo="1. DAS CONDIÇÕES GERAIS DA CONTRATAÇÃO", numero="1", elementos=[Elemento("p", "(Template interno) Detalhe completo da seção 1...")]),
        Secao(titulo="2. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO E FUNDAMENTAÇÃO LEGAL", numero="2", elementos=[Elemento("p", "(Template interno) Detalhe completo da seção 2...")]),
        Secao(titulo="3. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO CONSIDERADO O CICLO DE VIDA DO OBJETO E ESPECIFICAÇÃO DOS SERVIÇOS", numero="3", elementos=[Elemento("p", "(Template interno) Detalhe completo da seção 3...")]),
        Secao(titulo="4. REQUISITOS DA CONTRATAÇÃO", numero="4", elementos=[Elemento("p", "(Template interno) Detalhe completo da seção 4...")]),
        Secao(titulo="5. MODELO DE EXECUÇÃO CONTRATUAL", numero="5", elementos=[Elemento("p", "(Template interno) Detalhe completo da seção 5...")]),
        Secao(titulo="6. CRITÉRIOS DE MEDIÇÃO", numero="6", elementos=[Elemento("p", "(Template interno) Detalhe completo da seção 6...")]),
    ]

# ==========================
# Combinação de seções
# ==========================

def _index_por_numero(secoes: List[Secao]) -> Dict[str, Secao]:
    idx: Dict[str, Secao] = {}
    for s in secoes:
        chave = s.numero if s.numero is not None else s.titulo.strip().lower()
        idx[chave] = s
    return idx


def _ordenar_chaves_numericas(chaves: List[str]) -> List[str]:
    def keyfn(k: str):
        try:
            partes = [int(x) for x in k.split('.')]
            return partes
        except Exception:
            return [10**6]  # empurra não-numéricos para o fim
    return sorted(chaves, key=keyfn)


def combinar_secoes(
    secoes_modelo: List[Secao],
    secoes_template: List[Secao],
    modo: str = "complementar"
) -> List[Secao]:
    """Combina seções do DOCX de modelo com o template interno.

    modo:
      - 'modelo'       -> retorna somente `secoes_modelo`.
      - 'template'     -> retorna somente `secoes_template`.
      - 'complementar' -> base = modelo; se uma numeração do template **não existir** no modelo,
                          a seção do template é **inserida** (na ordem numérica).
    """
    if modo == "modelo":
        return secoes_modelo
    if modo == "template":
        return secoes_template

    # complementar
    idx_modelo = _index_por_numero(secoes_modelo)
    idx_template = _index_por_numero(secoes_template)

    # chaves numéricas presentes nas duas bases
    todas_chaves = set(idx_modelo.keys()) | set(idx_template.keys())
    chaves_ordenadas = _ordenar_chaves_numericas(list(todas_chaves))

    resultado: List[Secao] = []
    for chave in chaves_ordenadas:
        if chave in idx_modelo:
            resultado.append(idx_modelo[chave])
        else:
            # chave só existe no template -> inserir
            resultado.append(idx_template[chave])
    return resultado

# ==========================
# Geração do DOCX final
# ==========================

def _add_header_image_if_exists(doc: Document, path: Optional[str]):
    if not path:
        return
    try:
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        # Ajuste a largura conforme necessário
        run.add_picture(path, width=Inches(6.0))
    except Exception:
        pass


def _add_footer_image_if_exists(doc: Document, path: Optional[str]):
    if not path:
        return
    try:
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(path, width=Inches(6.0))
    except Exception:
        pass


def gerar_docx(
    secoes: List[Secao],
    caminho_saida: str,
    header_img: Optional[str] = None,
    footer_img: Optional[str] = None,
) -> str:
    """Gera um DOCX novo a partir das `secoes` combinadas.
    - Recria parágrafos como texto.
    - Para tabelas, replica o conteúdo (linhas/células) simples.
    Retorna o caminho do arquivo gerado.
    """
    doc = Document()

    _add_header_image_if_exists(doc, header_img)
    _add_footer_image_if_exists(doc, footer_img)

    for s in secoes:
        # Título da seção
        doc.add_paragraph(s.titulo)
        # Conteúdo
        for el in s.elementos:
            if el.tipo == "p":
                texto = str(el.payload).strip()
                if texto:
                    doc.add_paragraph(texto)
            elif el.tipo == "table":
                # Copiar estrutura básica da tabela
                try:
                    src_tbl = el.payload
                    rows = len(src_tbl.rows)
                    cols = len(src_tbl.columns)
                    if rows > 0 and cols > 0:
                        new_tbl = doc.add_table(rows=rows, cols=cols)
                        for r in range(rows):
                            for c in range(cols):
                                new_tbl.cell(r, c).text = src_tbl.cell(r, c).text
                except Exception:
                    # fallback: ignorar tabela se der erro
                    pass
        # espaço entre seções
        doc.add_paragraph("")

    doc.save(caminho_saida)
    return caminho_saida

# ==========================
# Exemplo de uso (Streamlit)
# ==========================

"""
# No seu app Streamlit:

import streamlit as st
from io import BytesIO

st.subheader("Monte o TR a partir de um modelo DOCX + template interno")
modelo = st.file_uploader("Envie o modelo DOCX (opcional)", type=["docx"])
modo = st.selectbox(
    "Como combinar?",
    ["complementar", "modelo", "template"],
    index=0,
    help=(
        "complementar: usa o conteúdo do DOCX e preenche seções ausentes com o template interno;\n"
        "modelo: usa só o DOCX;\n"
        "template: usa só o template interno."
    ),
)

if st.button("Gerar DOCX"):
    secoes_template = template_interno_padrao()
    secoes_modelo = []
    if modelo is not None:
        secoes_modelo = ler_modelo_docx(modelo)

    secoes_final = combinar_secoes(secoes_modelo, secoes_template, modo=modo)

    buffer = BytesIO()
    # ajuste os caminhos de cabeçalho/rodapé se desejar
    caminho_temp = "/mnt/data/TR_final.docx"
    gerar_docx(
        secoes_final,
        caminho_temp,
        header_img="/mnt/data/logo-prefeitura.png",
        footer_img="/mnt/data/rodapé.png",
    )

    with open(caminho_temp, "rb") as f:
        st.download_button(
            label="Baixar TR final (DOCX)",
            data=f.read(),
            file_name="TR_final.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
"""
