# Agente de Licitações – Gerador de Termo de Referência (TR)
# Versão interativa (wizard) com pré-visualização e exportação .docx
# Base legal: arts. 6º, 40 e 92 da Lei 14.133/2021 e art. 30 do Decreto Municipal nº 09/2024

import json
from datetime import date
from io import BytesIO

import streamlit as st
from docx import Document

# ==========================
# CONFIGURAÇÃO DA PÁGINA
# ==========================
st.set_page_config(
    page_title="Agente de Licitações – TR",
    layout="wide",
    page_icon="📑",
)

# Pequeno estilo para deixar parecido com apps web profissionais
st.markdown(
    """
    <style>
    .block-container {padding-top: 1.5rem; padding-bottom: 2rem;}
    .stProgress > div > div > div > div { background-color: #2E7D32; }
    .badge {background:#F1F5F9;border-radius:999px;padding:.2rem .6rem;font-size:.8rem;margin-left:.5rem}
    .muted{color:#64748B}
    textarea{font-family: ui-monospace, SFMono-Regular, Menlo, monospace;}
    </style>
    """,
    unsafe_allow_html=True,
)

# ==========================
# SIDEBAR – LOGO E NAVEGAÇÃO
# ==========================
st.sidebar.title("⚙️ Configuração")
logo = st.sidebar.file_uploader("Opcional: enviar logotipo (PNG/JPG)", type=["png", "jpg", "jpeg"]) 

passos = [
    "1) Objeto e categoria",
    "2) Parâmetros contratuais",
    "3) Requisitos",
    "4) Solução e ciclo de vida",
    "5) Critérios de medição",
    "6) Prévia e exportação",
]

etapa = st.sidebar.radio("Navegação", passos, index=0)

# Barra de progresso simples
st.sidebar.progress((passos.index(etapa) + 1) / len(passos))

# Salvar/Carregar rascunho
st.sidebar.markdown("---")
st.sidebar.subheader("💾 Rascunho")
rascunho_file = st.sidebar.file_uploader("Carregar rascunho (.json)", type=["json"], key="upjson")
if rascunho_file is not None:
    try:
        st.session_state["dados"] = json.load(rascunho_file)
        st.sidebar.success("Rascunho carregado!")
    except Exception as e:
        st.sidebar.error(f"Erro ao ler JSON: {e}")

# Inicializa estrutura de dados na sessão
if "dados" not in st.session_state:
    st.session_state["dados"] = {
        "objeto": "",
        "categoria": "Compra",
        "subcategoria": "Gás de cozinha (GLP)",
        "srp": True,
        "modalidade": "Pregão Eletrônico",
        "criterio": "Menor preço por item",
        "vigencia_meses": 12,
        "prorrogavel": True,
        "dotacao": "",
        "requisitos_marcados": [],
        "requisitos_extra": "",
        "solucao_texto": "",
        "ciclo_texto": "",
        "medicao_texto": "",
        "unidades_entrega": "",
    }

d = st.session_state["dados"]

# Utilitário: baixar rascunho
def baixar_json():
    return BytesIO(json.dumps(st.session_state["dados"], ensure_ascii=False, indent=2).encode("utf-8"))

st.sidebar.download_button(
    "⬇️ Baixar rascunho (.json)",
    data=baixar_json(),
    file_name="rascunho_tr.json",
    mime="application/json",
)

# ==========================
# CABEÇALHO VISUAL
# ==========================
col1, col2 = st.columns([1, 3])
with col1:
    if logo:
        st.image(logo, use_column_width=True)
    else:
        st.markdown("### Prefeitura de Brasnorte\n**Gerador de Termo de Referência**")
with col2:
    st.caption("Versão interativa. Preencha as etapas ao lado e gere o TR completo com base na Lei 14.133/2021.")

st.markdown("---")

# ==========================
# ETAPA 1 – OBJETO E CATEGORIA
# ==========================
if etapa == "1) Objeto e categoria":
    st.subheader("1) Objeto e categoria")
    d["objeto"] = st.text_area(
        "📝 Descreva o objeto da contratação (seja específico)",
        value=d.get("objeto", ""),
        placeholder=(
            "Ex.: Registro de preços para futura e eventual contratação de empresa especializada "
            "no fornecimento de gás de cozinha (GLP), botijões P13 e P45, com entrega nas unidades..."
        ),
        height=130,
    )

    c1, c2, c3 = st.columns(3)
    with c1:
        d["categoria"] = st.selectbox("Categoria do objeto", ["Compra", "Serviço", "Obra"], index=["Compra","Serviço","Obra"].index(d.get("categoria","Compra")))
    with c2:
        # Subcategorias dependem da categoria (exemplos)
        sub_opts = {
            "Compra": ["Gás de cozinha (GLP)", "Material de expediente", "Gêneros alimentícios", "Combustíveis"],
            "Serviço": ["Hotelaria e hospedagem", "Vigilância/segurança", "Limpeza e conservação", "Transporte"],
            "Obra": ["Reforma predial", "Pavimentação", "Construção civil"]
        }
        d["subcategoria"] = st.selectbox("Subcategoria", sub_opts[d["categoria"]], index=max(0, sub_opts[d["categoria"]].index(d.get("subcategoria", sub_opts[d["categoria"]][0])) if d.get("subcategoria") in sub_opts[d["categoria"]] else 0))
    with c3:
        d["unidades_entrega"] = st.text_input("Locais/Unidades atendidas (opcional)", value=d.get("unidades_entrega", ""), placeholder="Ex.: Escolas, CRAS, UBS, Almoxarifado central…")

    st.info("Dica: quanto mais específico o objeto, mais assertivo será o texto do TR gerado.")

# ==========================
# ETAPA 2 – PARÂMETROS CONTRATUAIS
# ==========================
if etapa == "2) Parâmetros contratuais":
    st.subheader("2) Parâmetros contratuais")
    c1, c2, c3 = st.columns(3)
    with c1:
        d["srp"] = st.checkbox("Usar Sistema de Registro de Preços (SRP)", value=d.get("srp", True))
    with c2:
        d["modalidade"] = st.selectbox("Modalidade", ["Pregão Eletrônico", "Concorrência", "Dispensa", "Inexigibilidade"], index=["Pregão Eletrônico","Concorrência","Dispensa","Inexigibilidade"].index(d.get("modalidade","Pregão Eletrônico")))
    with c3:
        d["criterio"] = st.selectbox("Critério de julgamento", ["Menor preço por item", "Menor preço global", "Técnica e preço"], index=["Menor preço por item","Menor preço global","Técnica e preço"].index(d.get("criterio","Menor preço por item")))

    c4, c5, c6 = st.columns(3)
    with c4:
        d["vigencia_meses"] = st.number_input("Vigência (meses)", min_value=1, max_value=60, value=int(d.get("vigencia_meses", 12)))
    with c5:
        d["prorrogavel"] = st.checkbox("Contrato prorrogável", value=d.get("prorrogavel", True))
    with c6:
        d["dotacao"] = st.text_input("Dotação (opcional)", value=d.get("dotacao", ""), placeholder="Ex.: 03.001.04.122.0002…")

# ==========================
# ETAPA 3 – REQUISITOS
# ==========================
if etapa == "3) Requisitos":
    st.subheader("3) Requisitos da contratação")

    # Requisitos padrão (checklist)
    req_base = [
        "Regularidade fiscal, previdenciária e trabalhista",
        "Capacidade técnica compatível (atestados)",
        "Atendimento a normas de segurança e qualidade",
        "Equipe técnica qualificada",
        "Plano de trabalho e cronograma",
        "Garantia de fornecimento ininterrupto",
        "Conformidade com LGPD e segurança da informação",
        "Responsabilidade socioambiental",
        "Observância às normas ABNT aplicáveis",
        "Apresentação de notas fiscais e relatórios para medição",
    ]

    # Extras por subcategoria (exemplos)
    extras_por_sub = {
        "Gás de cozinha (GLP)": [
            "Licenças e autorizações para comercialização e transporte de GLP",
            "Cilindros e botijões certificados pelo INMETRO",
            "Entrega em conformidade com normas de segurança (NRs, Corpo de Bombeiros)",
            "Peso líquido aferido e lacres de segurança",
        ],
        "Hotelaria e hospedagem": [
            "Recepção 24h e café da manhã incluso (quando aplicável)",
            "Wi-Fi, ar-condicionado, higiene e enxoval",
            "NF individual por hóspede e diária",
        ],
    }

    marcados = set(d.get("requisitos_marcados", []))

    st.markdown("**Selecione os requisitos aplicáveis**")
    # Base
    for r in req_base:
        val = st.checkbox(r, value=(r in marcados), key=f"req_{hash(r)}")
        if val:
            marcados.add(r)
        else:
            if r in marcados:
                marcados.remove(r)
    # Extras
    for r in extras_por_sub.get(d["subcategoria"], []):
        val = st.checkbox(r, value=(r in marcados), key=f"req_extra_{hash(r)}")
        if val:
            marcados.add(r)
        else:
            if r in marcados:
                marcados.remove(r)

    d["requisitos_marcados"] = list(marcados)

    d["requisitos_extra"] = st.text_area(
        "➕ Requisitos adicionais (texto livre)",
        value=d.get("requisitos_extra", ""),
        placeholder="Inclua exigências específicas do objeto/mercado local.",
    )

# ==========================
# ETAPA 4 – SOLUÇÃO E CICLO DE VIDA
# ==========================
if etapa == "4) Solução e ciclo de vida":
    st.subheader("4) Descrição da solução e ciclo de vida")

    # Sugestões automáticas (editáveis)
    sol1 = (
        "Solução 1 – Execução direta pela Administração: demanda estrutura, equipe e logística próprias; "
        "em regra, revela-se inviável por ausência de aparato técnico-operacional e pela economicidade."
    )
    sol2 = (
        "Solução 2 – Contratação de empresa especializada por licitação: garante expertise técnica, "
        "cumprimento de prazos, qualidade e conformidade regulatória."
    )
    conclusao = (
        "Conclusão: recomenda-se a execução indireta (Solução 2), por melhor custo-benefício, segurança jurídica "
        "e eficiência administrativa, em aderência à Lei 14.133/2021."
    )

    default_solucao = f"""{sol1}\n\n{sol2}\n\n{conclusao}"""

    d["solucao_texto"] = st.text_area(
        "✍️ Soluções avaliadas (edite se necessário)",
        value=d.get("solucao_texto", default_solucao),
        height=160,
    )

    default_ciclo = (
        "Planejamento da demanda → seleção do fornecedor → formalização (ARP/contrato) → execução e entregas → "
        "medições e pagamentos → avaliação de desempenho → encerramento contratual."
    )
    d["ciclo_texto"] = st.text_area(
        "♻️ Ciclo de vida do objeto (edite se necessário)",
        value=d.get("ciclo_texto", default_ciclo),
        height=120,
    )

# ==========================
# ETAPA 5 – CRITÉRIOS DE MEDIÇÃO
# ==========================
if etapa == "5) Critérios de medição":
    st.subheader("5) Critérios de medição")

    # Modelos por subcategoria (editáveis)
    med_glp = (
        "Unidade de medida: Kg de GLP e/ou diárias de botijões P13/P45 entregues.\n"
        "Itens mensurados: quantidade por tipo de botijão/cilindro; conformidade do peso líquido; lacres; \n"
        "documentos: NF com locais de entrega, datas/horários; comprovantes de recebimento; \n"
        "Aceite: conferência pelo fiscal; glosas em caso de divergências; pagamento apenas após aceite."
    )
    med_hotel = (
        "Unidade de medida: diária por hóspede (24h). Itens: tipo de acomodação, café da manhã, despesas autorizadas;\n"
        "Documentos: relação nominal (CPF, check-in/out), relatório mensal, NFs compatíveis com ARP/contrato;\n"
        "Aceite e atesto pelo gestor; glosas se houver inconsistência."
    )
    med_generico = (
        "Medição baseada em entregas/serviços efetivamente realizados, conferidos e atestados pelo fiscal; \n"
        "Pagamento condicionado ao aceite formal e conformidade com quantitativos, prazos e especificações."
    )

    sugestao = med_generico
    if d["subcategoria"] == "Gás de cozinha (GLP)":
        sugestao = med_glp
    elif d["subcategoria"] == "Hotelaria e hospedagem":
        sugestao = med_hotel

    d["medicao_texto"] = st.text_area(
        "🧾 Texto de medição (edite se necessário)",
        value=d.get("medicao_texto", sugestao),
        height=170,
    )

# ==========================
# FUNÇÃO – GERAR TEXTO DO TR
# ==========================

def gerar_texto_tr(d: dict) -> str:
    obj = d.get("objeto", "").strip()
    sub = d.get("subcategoria", "")

    # Cabeçalho institucional
    partes = [
        "PREFEITURA MUNICIPAL DE BRASNORTE - MT",
        "SECRETARIA MUNICIPAL DE ADMINISTRAÇÃO",
        "",
        "TERMO DE REFERÊNCIA",
        "",
        "1. DAS CONDIÇÕES GERAIS DA CONTRATAÇÃO",
        f"\n1.1 O presente termo de referência tem por objeto o {obj.upper() if obj else '[OBJETO NÃO INFORMADO]'}, com sede localizada no município de Brasnorte-MT, em conformidade com as especificações de descrição e quantidade detalhadamente elencadas neste documento, amparada pelas disposições legais vigentes que regulam tal procedimento, visando atender as necessidades da Prefeitura Municipal de Brasnorte-MT e de suas Secretarias Municipais;",
        "\n1.2 O objeto desta contratação não se enquadra como sendo de bem de luxo, conforme Decreto Municipal nº 03/2024;",
        f"\n1.3 O prazo de vigência da contratação é de {d.get('vigencia_meses',12)} (doze) meses, contados da data de assinatura da ARP (Ata de Registro de Preços) ou do Contrato, na forma do art. 105 da Lei nº 14.133/2021, podendo ser prorrogado a critério da Administração Pública.",
        ("\n1.4 O prazo de vigência poderá ser prorrogado, desde que haja interesse de ambas as partes, nos termos dos arts. 106 e 107 da Lei nº 14.133/2021." if d.get("prorrogavel", True) else ""),
        "",
        "2. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO E FUNDAMENTAÇÃO LEGAL",
        f"\n2.1 A presente contratação se fundamenta na necessidade institucional de garantir o atendimento contínuo e eficiente relativo a {sub.lower()} para as unidades/secretarias do Município, assegurando a continuidade dos serviços públicos e o cumprimento da missão institucional.",
        "\n2.2 A contratação tem por finalidade atender às Secretarias Municipais, promovendo suporte às atividades administrativas, operacionais e técnicas, conforme planejamento da gestão e demanda das unidades requisitantes;",
        "\n2.3 Justifica-se pela inexistência de estrutura própria que permita a execução direta, sendo necessária a contratação especializada para garantir eficiência, economicidade e regularidade;",
        f"\n2.4 O procedimento adotado será a modalidade {d.get('modalidade')} com critério de julgamento {d.get('criterio').lower()}, observado o disposto na Lei nº 14.133/2021." ,
        ("\n2.5 Será adotado o Sistema de Registro de Preços (SRP), nos termos do Decreto Federal nº 11.462/2023, proporcionando flexibilidade e economicidade." if d.get("srp", True) else ""),
        "",
        "3. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO CONSIDERADO O CICLO DE VIDA DO OBJETO E ESPECIFICAÇÃO",
        f"\n3.1 {d.get('solucao_texto','')}",
        f"\n3.2 Ciclo de vida do objeto: {d.get('ciclo_texto','')} ",
        "",
        "4. REQUISITOS DA CONTRATAÇÃO",
        "\n4.1 A contratada deverá atender, no mínimo, aos seguintes requisitos:" ,
    ]

    # Lista de requisitos marcados
    for i, r in enumerate(d.get("requisitos_marcados", []), start=1):
        partes.append(f"\n4.1.{i} {r}.")
    if d.get("requisitos_extra"):
        partes.append(f"\n4.2 Requisitos adicionais: {d['requisitos_extra']}")

    partes += [
        "",
        "5. MODELO DE EXECUÇÃO CONTRATUAL",
        "\n5.1 A execução contratual se dará mediante ordens de fornecimento/serviço emitidas pela Administração, com fiscalização designada conforme a Lei nº 14.133/2021;",
        "\n5.2 Os pagamentos ocorrerão após aceite formal, mediante apresentação de nota fiscal e relatório de entrega/execução, em conformidade com as especificações e quantitativos contratados;",
        "\n5.3 A gestão e a fiscalização observarão os arts. 117 a 124 da Lei nº 14.133/2021, incluindo registros de ocorrências, notificações e relatórios de acompanhamento;",
        "",
        "6. CRITÉRIOS DE MEDIÇÃO",
        f"\n6.1 {d.get('medicao_texto','Medição baseada em entregas/serviços atestados pelo fiscal, com pagamento após aceite.')}",
        "",
        f"**Brasnorte - MT, {date.today().strftime('%d/%m/%Y')}**",
        "\n---\n",
        "Este documento é gerado automaticamente com base nas diretrizes legais vigentes e poderá ser personalizado conforme peculiaridades do objeto. Recomenda-se revisão da Procuradoria Jurídica e do Controle Interno.",
    ]

    return "\n".join([p for p in partes if p is not None and p != ""])

# ==========================
# ETAPA 6 – PRÉVIA E EXPORTAÇÃO
# ==========================
if etapa == "6) Prévia e exportação":
    st.subheader("6) Pré-visualização e exportação")

    # Validações simples
    if not d.get("objeto"):
        st.error("Informe o objeto da contratação na etapa 1.")
    else:
        texto = gerar_texto_tr(d)
        st.markdown("### 📄 Prévia do Termo de Referência")
        st.text_area("Conteúdo gerado:", texto, height=500)

        # Exportar DOCX
        def gerar_docx(texto: str) -> BytesIO:
            doc = Document()
            if logo:
                # Insere logotipo no topo (se o usuário enviou). Tamanho dependerá da imagem.
                doc.add_picture(logo)
            for linha in texto.split("\n"):
                doc.add_paragraph(linha)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        docx_bin = gerar_docx(texto)
        st.download_button(
            label="📥 Baixar TR em Word (.docx)",
            data=docx_bin,
            file_name="Termo_de_Referencia_Brasnorte.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Exportar rascunho JSON desta tela
        st.download_button(
            label="💾 Baixar rascunho (.json)",
            data=baixar_json(),
            file_name="rascunho_tr.json",
            mime="application/json",
        )

# Rodapé
st.markdown("<span class='muted'>Agente de Licitações – Prefeitura de Brasnorte • Lei 14.133/2021</span>", unsafe_allow_html=True)
